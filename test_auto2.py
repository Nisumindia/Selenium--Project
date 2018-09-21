
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.support.ui import Select
from docx import Document
from docx.shared import Inches
import time
import random

document = Document()
p = document.add_paragraph()
r = p.add_run()


driver = webdriver.Chrome()
driver.get("https://www.mcom-061.tbe.zeus.fds.com")

driver.maximize_window()

driver.save_screenshot("/Users/premkumer/Downloads/1-homewindow.png")
r.add_picture("/Users/premkumer/Downloads/1-homewindow.png",width=Inches(8.0), height=Inches(5.0))

elem = driver.find_element_by_id("shopByDepartmentLabelText")
elem.click()
driver.implicitly_wait(10)
driver.save_screenshot("/Users/premkumer/Downloads/2-flyouts.png")
r.add_picture("/Users/premkumer/Downloads/2-flyouts.png",width=Inches(8.0), height=Inches(5.0))

element = driver.find_element_by_link_text("HOME")
hover = ActionChains(driver).move_to_element(element)
hover.perform()
elem = driver.find_element_by_link_text("Bakeware")
elem.click()
driver.save_screenshot("/Users/premkumer/Downloads/3-catalog.png")
r.add_picture("/Users/premkumer/Downloads/3-catalog.png",width=Inches(8.0), height=Inches(5.0))


foo = ['22805','1225', '298391','img_2987176']
secure_random = random.SystemRandom()
Item=secure_random.choice(foo)

driver.find_element_by_id("globalSearchInputField").clear()
driver.find_element_by_id("globalSearchInputField").send_keys(Item)
driver.find_element_by_id("searchSubmit").click()
driver.implicitly_wait(10)
driver.save_screenshot("/Users/premkumer/Downloads/4-item.png")
r.add_picture("/Users/premkumer/Downloads/4-item.png",width=Inches(8.0), height=Inches(5.0))
driver.find_element_by_id("bag-add-"+Item).click()
driver.implicitly_wait(10)
driver.save_screenshot("/Users/premkumer/Downloads/5-add-to-bag.png")
r.add_picture("/Users/premkumer/Downloads/5-add-to-bag.png",width=Inches(8.0), height=Inches(5.0))
elems = driver.find_elements_by_link_text("VIEW BAG & CHECKOUT")
if len(elems) > 0 :
 driver.find_element_by_link_text("VIEW BAG & CHECKOUT").click()
else:
 driver.find_element_by_id("bag-add-"+Item).click()
 driver.find_element_by_link_text("VIEW BAG & CHECKOUT").click()



driver.save_screenshot("/Users/premkumer/Downloads/6-view&checkout.png")
r.add_picture("/Users/premkumer/Downloads/6-view&checkout.png",width=Inches(8.0), height=Inches(5.0))

elems = driver.find_elements_by_id("continueCheckout")
if len(elems) > 0 :
 driver.find_element_by_id("continueCheckout").click()
 time.sleep(10) 
else:
 driver.find_element_by_id("CHECKOUT").click()
driver.save_screenshot("/Users/premkumer/Downloads/7-checkout.png")
r.add_picture("/Users/premkumer/Downloads/7-checkout.png",width=Inches(8.0), height=Inches(5.0))

time.sleep(10)
driver.find_element_by_id("guest-checkout").click()
driver.implicitly_wait(10)
driver.save_screenshot("/Users/premkumer/Downloads/8-guest-checkout.png")
r.add_picture("/Users/premkumer/Downloads/8-guest-checkout.png",width=Inches(8.0), height=Inches(5.0))
driver.find_element_by_id("rc-shipping-firstName").clear()
driver.find_element_by_id("rc-shipping-firstName").send_keys("Jeniya")
driver.find_element_by_id("rc-shipping-lastName").clear()
driver.find_element_by_id("rc-shipping-lastName").send_keys("Bechtelar")
driver.find_element_by_id("rc-shipping-line1").clear()
driver.find_element_by_id("rc-shipping-line1").send_keys("602 W Dryden St")
driver.implicitly_wait(10)
driver.find_element_by_id("rc-shipping-postal-code").clear()
driver.find_element_by_id("rc-shipping-postal-code").send_keys("91202")
time.sleep(20)
driver.find_element_by_id("rc-shipping-state")
driver.find_element_by_id("rc-shipping-phone").clear()
driver.find_element_by_id("rc-shipping-phone").send_keys("(233) 243-4234")
driver.find_element_by_id("rc-shipping-continue").click()
time.sleep(10)
elems = driver.find_elements_by_id("rc-shipping-group-continue")
if len(elems) > 0 :
 driver.find_element_by_id("rc-shipping-group-continue").click()
 time.sleep(10) 
driver.find_element_by_id("rc-payment-card-number").clear()
driver.find_element_by_id("rc-payment-card-number").send_keys("4445222299990007")
Select(driver.find_element_by_id("rc-payment-card-type")).select_by_visible_text("Visa")
time.sleep(10)
Select(driver.find_element_by_id("rc-payment-card-month")).select_by_visible_text("03")
Select(driver.find_element_by_id("rc-payment-card-year")).select_by_visible_text("2022")
driver.find_element_by_id("rc-payment-scode").clear()
driver.find_element_by_id("rc-payment-scode").send_keys("222")
driver.find_element_by_id("rc-payment-email").clear()
driver.find_element_by_id("rc-payment-email").send_keys("test@macys.com")
driver.find_element_by_id("rc-payment-continue").click()
time.sleep(10)
driver.find_element_by_id("rc-place-order").click()
driver.save_screenshot("/Users/premkumer/Downloads/8-place_order.png")
r.add_picture("/Users/premkumer/Downloads/8-place_order.png",width=Inches(8.0), height=Inches(5.0))
time.sleep(10)
driver.save_screenshot("/Users/premkumer/Downloads/9-thank_you.png")
r.add_picture("/Users/premkumer/Downloads/9-thank_you.png",width=Inches(8.0), height=Inches(5.0))
document.save('demo.docx')
#driver.close()
#driver.find_element_by_id("guest-checkout").click()


